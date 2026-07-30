"""Apply a sparse block-edit diff back into the original .docx's runs, and
build the side-by-side preview JSON (spec.md §8, architecture.md §6).

Only run text is ever touched — run.text's setter rewrites the run's
<w:t> content without touching its rPr (font/size/color/bold/etc.), so
formatting on both edited and untouched runs is preserved by construction.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from app.services.docx_blocks import Block, BlockLocation
from app.services.llm_rewrite import BlockEdit


def _get_runs(document: Document, location: BlockLocation) -> list:
    if location.kind == "paragraph":
        paragraph = document.paragraphs[location.paragraph_index]
    else:
        table = document.tables[location.table_index]
        cell = table.rows[location.row_index].cells[location.cell_index]
        paragraph = cell.paragraphs[location.paragraph_index]
    # run_indices is empty on blocks persisted before it existed.
    indices = location.run_indices or [location.run_index]
    return [paragraph.runs[i] for i in indices]


def apply_edits_to_docx(source_path: Path, blocks: list[Block], edits: list[BlockEdit], output_path: Path) -> None:
    edits_map = {e.id: e.text for e in edits}
    document = Document(str(source_path))

    for block in blocks:
        new_text = edits_map.get(block.id)
        if new_text is None:
            continue
        runs = _get_runs(document, block.location)
        # A block may cover several original runs (merged mid-word/mid-
        # sentence splits, see docx_blocks._group_runs); the edited text
        # goes into the first run's original location, the rest are
        # cleared so the merged block doesn't duplicate text.
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def _group_key(location: BlockLocation) -> str:
    """Identifies which paragraph a run belongs to, so the frontend can
    render runs as flowing paragraphs instead of one row per text run."""
    if location.kind == "paragraph":
        return f"p{location.paragraph_index}"
    return f"t{location.table_index}-row{location.row_index}-c{location.cell_index}-p{location.paragraph_index}"


def build_preview(blocks: list[Block], edits: list[BlockEdit]) -> list[dict]:
    edits_map = {e.id: e.text for e in edits}
    preview = []
    for block in blocks:
        changed = block.id in edits_map
        preview.append(
            {
                "id": block.id,
                "group_key": _group_key(block.location),
                "original_text": block.text,
                "tailored_text": edits_map.get(block.id, block.text),
                "style_name": block.style_name,
                "run_style": block.run_style.model_dump(),
                "section": block.section,
                "editable": block.editable,
                "changed": changed,
            }
        )
    return preview


def tailored_text(blocks: list[Block], edits: list[BlockEdit]) -> str:
    edits_map = {e.id: e.text for e in edits}
    return " ".join(edits_map.get(b.id, b.text) for b in blocks)


def _slugify_first_line(text: str) -> str:
    first_line = next((line.strip() for line in text.splitlines() if line.strip()), "")
    slug = re.sub(r"[^a-z0-9]+", "-", first_line.lower()).strip("-")
    return slug[:40] or "role"


def make_output_filename(original_filename: str, job_description: str) -> str:
    stem = Path(original_filename).stem
    slug = _slugify_first_line(job_description)
    return f"{stem}-tailored-{slug}.docx"
