"""Job description ingestion: pasted text, or text extracted from a
.docx/.pdf upload. Formatting is irrelevant here — only the text is ever
used (spec.md §5.1) — so extraction just needs to be readable, not
formatting-faithful.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

JOB_DESCRIPTION_FILENAME = "job_description.txt"


def extract_text_from_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
    return "\n".join(parts)


def extract_text_from_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(p for p in parts if p.strip())


def save_job_description(session_dir: Path, text: str) -> None:
    (session_dir / JOB_DESCRIPTION_FILENAME).write_text(text)


def load_job_description(session_dir: Path) -> str | None:
    path = session_dir / JOB_DESCRIPTION_FILENAME
    if not path.exists():
        return None
    return path.read_text()
