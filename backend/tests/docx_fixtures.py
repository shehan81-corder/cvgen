"""Builds sample .docx files in-memory (via python-docx) for tests.

Generating fixtures with python-docx, rather than committing binary .docx
files, keeps them small, readable in diffs, and easy to vary per test.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


def build_simple_cv(path: Path) -> Path:
    document = Document()

    heading = document.add_paragraph("JOHN SMITH")
    heading.style = document.styles["Heading 1"]

    contact = document.add_paragraph()
    run = contact.add_run("john.smith@example.com | +1 555-123-4567")
    run.font.size = Pt(10)

    summary_heading = document.add_paragraph("Summary")
    summary_heading.style = document.styles["Heading 2"]

    summary = document.add_paragraph()
    summary_run = summary.add_run(
        "Backend engineer with 8 years building payment systems."
    )
    summary_run.font.name = "Calibri"
    summary_run.font.size = Pt(11)
    summary_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    exp_heading = document.add_paragraph("Experience")
    exp_heading.style = document.styles["Heading 2"]

    title_line = document.add_paragraph()
    title_run = title_line.add_run("Senior Engineer, Acme Corp")
    title_run.bold = True

    date_line = document.add_paragraph()
    date_line.add_run("Jan 2020 - Present")

    bullet = document.add_paragraph(style="List Bullet")
    bullet_run = bullet.add_run(
        "Led migration of the billing service to a new provider."
    )
    bullet_run.italic = True

    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Python")
    table.rows[0].cells[1].paragraphs[0].add_run("PostgreSQL")

    document.save(str(path))
    return path


def build_cv_with_skills_table(path: Path) -> Path:
    document = Document()

    document.add_paragraph("JANE DOE").style = document.styles["Heading 1"]

    skills_heading = document.add_paragraph("Skills")
    skills_heading.style = document.styles["Heading 2"]

    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Languages")
    table.rows[0].cells[1].paragraphs[0].add_run("Python, Go")
    table.rows[1].cells[0].paragraphs[0].add_run("Cloud")
    table.rows[1].cells[1].paragraphs[0].add_run("AWS, GCP")

    education_heading = document.add_paragraph("Education")
    education_heading.style = document.styles["Heading 2"]

    degree_line = document.add_paragraph()
    degree_line.add_run("BSc Computer Science, State University")

    document.save(str(path))
    return path


def build_cv_no_heading_styles(path: Path) -> Path:
    """A CV that marks sections with bold all-caps text instead of Heading styles."""
    document = Document()

    document.add_paragraph("ALEX RIVERA")

    section = document.add_paragraph()
    section_run = section.add_run("EXPERIENCE")
    section_run.bold = True

    title_line = document.add_paragraph()
    title_line.add_run("Product Manager, Globex Inc")

    date_line = document.add_paragraph()
    date_line.add_run("Mar 2019 - Dec 2022")

    bullet = document.add_paragraph(style="List Bullet")
    bullet.add_run("Managed a team of five designers and engineers.")

    document.save(str(path))
    return path
