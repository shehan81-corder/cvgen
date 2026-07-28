from docx import Document

from tests.pdf_fixtures import build_minimal_pdf

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _build_jd_docx(path) -> bytes:
    document = Document()
    document.add_paragraph("Senior Backend Engineer")
    document.add_paragraph("We need someone strong in Python and distributed systems.")
    document.save(str(path))
    return path.read_bytes()


def test_paste_text_job_description_stored_and_retrievable(client, session_id):
    text = "Senior Backend Engineer role requiring Python and AWS experience."
    post_response = client.post(
        f"/sessions/{session_id}/job-description", data={"text": text}
    )
    assert post_response.status_code == 200

    get_response = client.get(f"/sessions/{session_id}/job-description")
    assert get_response.status_code == 200
    assert get_response.json()["text"] == text


def test_upload_docx_job_description_extracts_text(client, session_id, tmp_path):
    content = _build_jd_docx(tmp_path / "jd.docx")
    response = client.post(
        f"/sessions/{session_id}/job-description",
        files={"file": ("jd.docx", content, DOCX_CONTENT_TYPE)},
    )
    assert response.status_code == 200

    stored = client.get(f"/sessions/{session_id}/job-description").json()["text"]
    assert "Senior Backend Engineer" in stored
    assert "Python and distributed systems" in stored


def test_upload_pdf_job_description_extracts_text(client, session_id):
    pdf_bytes = build_minimal_pdf("Senior Backend Engineer needing Python skills")
    response = client.post(
        f"/sessions/{session_id}/job-description",
        files={"file": ("jd.pdf", pdf_bytes, "application/pdf")},
    )
    assert response.status_code == 200

    stored = client.get(f"/sessions/{session_id}/job-description").json()["text"]
    assert "Senior Backend Engineer" in stored
    assert "Python skills" in stored


def test_job_description_requires_session(client):
    response = client.post(
        "/sessions/does-not-exist/job-description", data={"text": "hello"}
    )
    assert response.status_code == 404


def test_job_description_requires_text_or_file(client, session_id):
    response = client.post(f"/sessions/{session_id}/job-description")
    assert response.status_code == 400


def test_job_description_rejects_both_text_and_file(client, session_id, tmp_path):
    content = _build_jd_docx(tmp_path / "jd.docx")
    response = client.post(
        f"/sessions/{session_id}/job-description",
        data={"text": "some text"},
        files={"file": ("jd.docx", content, DOCX_CONTENT_TYPE)},
    )
    assert response.status_code == 400


def test_get_job_description_before_upload_404(client, session_id):
    response = client.get(f"/sessions/{session_id}/job-description")
    assert response.status_code == 404
